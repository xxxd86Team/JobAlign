import streamlit as st
import pandas as pd
import json
import plotly.express as px
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import time
from openai import OpenAI
from PIL import Image
import pytesseract

# ================= 1. 全局配置与状态管理 =================
st.set_page_config(
    page_title="JobAlign AI Pro",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 初始化 Session State
if 'analyzed' not in st.session_state:
    st.session_state.analyzed = False
if 'result_json' not in st.session_state:
    st.session_state.result_json = None
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""


# ================= 2. 核心处理工具类 =================

class DocumentHandler:
    @staticmethod
    def extract_text(file):
        """
        统一处理 PDF / Word / 文本 / 图片 的文本提取
        支持：
        - .pdf
        - .doc / .docx
        - .txt
        - 图片：.png / .jpg / .jpeg / .bmp / .tiff / .gif（通过 OCR 识别）
        """
        text = ""
        try:
            filename = getattr(file, "name", "")
            ext = filename.split(".")[-1].lower() if "." in filename else ""

            # 确保指针在文件开头
            try:
                file.seek(0)
            except Exception:
                pass

            if ext == 'pdf':
                reader = PdfReader(file)
                for page in reader.pages:
                    content = page.extract_text()
                    if content:
                        text += content + "\n"

            elif ext in ['docx', 'doc']:
                doc = Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])

            elif ext == 'txt':
                text = file.getvalue().decode("utf-8")

            elif ext in ['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif']:
                # 图片 OCR 识别
                file_bytes = file.read()
                image = Image.open(io.BytesIO(file_bytes))
                # 如本机有中文语言包，可使用 lang='chi_sim+eng'
                text = pytesseract.image_to_string(image,lang='chi_sim+eng')

            else:
                # 兜底：尝试文本方式读取
                try:
                    text = file.getvalue().decode("utf-8")
                except Exception:
                    text = ""

            return text
        except Exception as e:
            return f"Error: 文件解析失败 ({str(e)})"


class WordGenerator:
    @staticmethod
    def create_docx_from_markdown(markdown_text):
        """将 Markdown 格式的简历草稿转换为格式化的 Word 文档"""
        doc = Document()

        # 设置基础样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        lines = markdown_text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 处理一级标题 (# Title)
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 处理二级标题 (## Title)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)

            # 处理三级标题 (### Title)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)

            # 处理列表项 (- Item / * Item)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')

            # 处理加粗 (**Text**) - 简单处理，仅去除标记
            elif '**' in line:
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)

            # 普通文本
            else:
                doc.add_paragraph(line)

        # 保存到内存流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# ================= 3. AI 交互逻辑 =================

MOCK_DATA = {
    "total_score": 78,
    "dimensions": {
        "技能匹配度": 82,
        "经验相关性": 75,
        "行业契合度": 70,
        "表达与亮点": 88
    },
    "highlights": [
        "演示模式：已有 AI/大模型相关项目，和 AI 产品/智能分析岗位高度相关。",
        "演示模式：具备一定 Python / 数据分析基础，方便后续向数据产品或智能分析方向延展。"
    ],
    "gaps": [
        "演示模式：简历中缺少系统性的指标设计与业务结果量化描述。",
        "演示模式：缺少对协作方式、跨部门沟通的具体案例说明。"
    ],
    "suggestions": [
        {
            "section": "项目经历",
            "original": "参与风向监控 Agent 项目。",
            "problem": "描述过于笼统，看不出业务背景、你的职责和结果。",
            "rewrite": "主导风向监控 Agent 需求分析与PRD撰写，覆盖战略/行业/产品等6个维度的信息源，将竞品情报整理耗时从1天缩短至2小时，为老板周例会提供结构化对手情报输入。"
        },
        {
            "section": "技能",
            "original": "熟练使用 Office。",
            "problem": "表述过泛，与目标岗位的关键能力缺乏关联。",
            "rewrite": "熟练使用 Excel / Pandas 进行数据清洗与漏斗分析，具备基础 SQL 查询能力，可独立完成简历数据与业务日志数据的结构化处理。"
        }
    ],
    "draft_resume": """# 演示简历
## 个人简介
我正在向 AI 产品 / 数据分析方向发展，具备基础产品方法论和数据分析能力，已经通过多项项目实践熟悉「需求分析 → 方案设计 → 项目落地 → 结果复盘」的完整闭环。

## 教育经历
- **某某大学 本科 · 专业：XXX**（201X - 202X）
  - 相关课程：数据分析、统计学基础、数据库原理、计算机基础

## 实习 / 项目经历
- **风向监控 Agent 项目｜AI 产品实习 / 个人项目**
  - 背景：公司缺乏系统的竞品与行业信息收集机制，情报依赖人工搜索与零散记录。
  - 职责：主导风向监控 Agent 的需求分析、PRD撰写与核心流程设计。
  - 方案：
    - 设计战略 / 行业 / 财务 / 产品 / 招投标 / 口碑 6个维度的监控框架；
    - 定义多 Agent 协同检索与聚合规则，搭建舆情与情报数据的初步结构化方案。
  - 结果：
    - 将竞品情报整理耗时从1天缩短至2小时；
    - 支持老板在周例会中更系统地评估竞对策略与行业动向。
  - 工具：Coze / 腾讯元宝、多Agent编排、Notion、Excel

- **JobAlign AI 简历匹配与成长规划工具｜个人项目**
  - 背景：求职者很难理解 JD 需求、评估自身匹配度并规划下一步学习。
  - 职责：从0到1设计并实现简历上传解析、JD对比、匹配报告、学习规划与岗位推荐的完整流程。
  - 方案：
    - 使用 Streamlit 构建前端界面，支持 PDF / Word / 图片简历与多个 JD 同时上传；
    - 利用大模型生成匹配度评分、亮点 / 缺失分析、简历改写建议与3-6个月成长路径。
  - 结果：
    - 帮助用户快速识别与自身背景更契合的岗位方向；
    - 提供结构化的学习与项目实践建议，为后续求职打基础。
  - 工具：Python、Streamlit、OpenAI / DeepSeek API、PyPDF2、pytesseract

## 技能
- 产品：需求分析、PRD撰写、用户场景拆解、UAT测试
- 数据：Python（Pandas）、SQL基础、简单可视化与漏斗分析
- 工具：Figma / 墨刀、Notion、Excel、Streamlit
- AI：对大模型、RAG、多 Agent 有一定理解，能基于平台搭建简单智能体流程

## 其他
- 持续在 CSDN / 个人公众号输出数据与 AI 产品相关内容，保持自我迭代。
""",
    "learning_plan": {
        "target_direction": "AI 产品 / 数据产品 方向",
        "summary": "综合你的简历与目标 JD，更推荐你在未来3-6个月重点强化：产品方法论、数据分析思维与项目复盘能力，用少量但高质量的项目支撑简历，而不是盲目堆数量。",
        "skills_to_focus": [
            "系统化的PRD写作与需求拆解",
            "SQL + 基础数据分析思维",
            "业务指标设计与结果量化表达",
            "项目复盘与结构化表达"
        ],
        "stages": [
            {
                "name": "第1-4周：打基础（理解岗位 & 强化表达）",
                "goals": [
                    "搞清楚AI产品 / 数据产品岗位的日常与核心能力",
                    "能写出结构清晰且有重点的PRD / 项目说明"
                ],
                "actions": [
                    "每周阅读2-3篇 AI / 数据产品案例拆解，将核心结构和亮点记录下来。",
                    "选1-2个你常用的产品，尝试从“问题-目标-方案-指标”的角度各写1页分析。"
                ]
            },
            {
                "name": "第5-8周：打造 1-2 个可写进简历的项目",
                "goals": [
                    "产出至少1个完整项目，可在简历中用1/3页重点描述",
                    "项目说明中能体现“做了什么”和“带来了什么变化”"
                ],
                "actions": [
                    "基于 JobAlign 等现有项目，补上需求背景、目标用户、关键指标与复盘思考。",
                    "结合公开数据或模拟数据，做一份简单的数据分析或看板，并写成小报告。"
                ]
            },
            {
                "name": "第9-12周：校准简历 & 预热面试",
                "goals": [
                    "让简历与目标 JD 的关键词高度对齐，同时保持真实",
                    "提前熟悉常见面试问法和项目深挖角度"
                ],
                "actions": [
                    "针对 3-5 条目标 JD，使用本工具多次优化简历表述，形成 1-2 份主力版本。",
                    "在牛客等平台刷同岗位面经，整理高频问题，并用自己的项目练习回答。"
                ]
            }
        ]
    },
    "resources": [
        {
            "platform": "B站",
            "category": "学习视频",
            "search_keyword": "产品经理 PRD 入门 案例 拆解",
            "reason": "帮助你系统理解 PRD 的结构和写法，提升简历中产品项目的专业度。"
        },
        {
            "platform": "B站",
            "category": "学习视频",
            "search_keyword": "Python SQL 数据分析 零基础 实战 项目",
            "reason": "你对数据分析有兴趣，但缺少成体系项目，可以通过实战教学补齐。"
        },
        {
            "platform": "牛客",
            "category": "面试经验",
            "search_keyword": "AI 产品 实习 面经 2024",
            "reason": "目标 JD 是 AI / 大模型相关产品方向，提前熟悉常见面试问题和考察维度。"
        },
        {
            "platform": "CSDN",
            "category": "技术文章",
            "search_keyword": "Streamlit 简历分析 项目 实战",
            "reason": "你已经在做 Streamlit 简历分析工具，可以参考他人实践，丰富项目亮点。"
        }
    ],
    "job_recommendations": [
        {
            "title": "AI 产品实习生",
            "company_type": "头部/新锐互联网公司（示例）",
            "location": "一线 / 新一线城市",
            "similarity_to_target_jd": 90,
            "match_reason": "岗位同样聚焦大模型 / 智能体方向，要求你具备产品思维与基础技术理解，与现有项目非常契合。",
            "core_requirements": [
                "参与 AI 产品需求分析、方案设计与文档撰写",
                "对主流大模型 / Agent 应用有基本了解，有实践经验更佳",
                "良好的沟通协作能力，能在技术与业务之间做有效对接"
            ]
        },
        {
            "title": "数据产品实习生",
            "company_type": "数据智能 / 企业服务公司（示例）",
            "location": "北上广深 / 杭州 / 成都",
            "similarity_to_target_jd": 85,
            "match_reason": "在保持产品岗位属性的前提下，更强调数据分析与指标设计，与你的 Python / SQL 和项目经历匹配度较高。",
            "core_requirements": [
                "参与数据产品需求梳理与指标体系设计",
                "配合中台 /业务方搭建分析报表与看板",
                "具备基础 SQL / Python 数据处理能力"
            ]
        }
    ],
    "target_jd_overview": [
        {
            "jd_index": 1,
            "jd_title": "JD_1：AI 产品实习生（示例）",
            "match_score": 88,
            "recommendation_level": "强烈推荐",
            "short_comment": "岗位方向与简历中的 AI 产品 / 智能体项目高度一致，是当前背景下优先级最高的选择之一。"
        },
        {
            "jd_index": 2,
            "jd_title": "JD_2：数据分析实习生（示例）",
            "match_score": 80,
            "recommendation_level": "可重点考虑",
            "short_comment": "强调数据分析能力和 SQL / Python，对你现在的技术基础比较友好，但产品成分略弱。"
        }
    ],
    "selected_jd_index": 1
}


def analyze_with_llm(api_key, base_url, model, resume, jd_list):
    """
    resume: 简历文本
    jd_list: [{'index': int, 'title': str, 'text': str}, ...]  支持多个 JD
    """
    client = OpenAI(api_key=api_key, base_url=base_url)

    # 组合多 JD 内容
    jd_blocks = []
    for idx, jd in enumerate(jd_list, start=1):
        title = jd.get("title", f"JD_{idx}")
        text = jd.get("text", "")
        jd_blocks.append(
            f"<<<JD_{idx} - {title}>>>\n{text[:2500]}"
        )
    jd_combined = "\n\n".join(jd_blocks)

    system_prompt = """
你是一名非常专业的「简历评估 + 职业发展教练」，熟悉校招 / 实习 / 社招 ATS 筛选逻辑，
理解 AI 产品 / 数据分析 / 互联网业务岗位的真实工作内容和用人标准。

你的目标：
- 帮求职者看清「当前简历」与「多个候选 JD」的匹配情况；
- 帮他选出更值得重点冲刺的岗位方向（不替他决定人生，只做专业建议）；
- 在此基础上，给出简历优化建议、未来3–6个月的成长规划、学习资源推荐，以及同方向的其他公司岗位参考。

请根据【简历】和【候选 JD 列表】进行分析，并返回严格的 JSON，字段必须包含：

1. total_score        (0-100整数)
   - 对「最终选中的 JD」的总体匹配度评分。

2. dimensions         (对象，键包括：
                        - 技能匹配度
                        - 经验相关性
                        - 行业契合度
                        - 表达与亮点
                       值为0-100整数)

3. highlights         (数组，3-5条高匹配点，每条为字符串，语言专业、具体，避免空洞鸡汤)

4. gaps               (数组，3-5条缺失或风险点，每条为字符串，尽量关联到面试 / ATS 筛选风险)

5. suggestions        (数组，元素为对象，字段：
                       - section: 所属模块，如“项目经历”“实习经历”“技能”
                       - original: 简历原文句子
                       - problem: 存在的问题（例如：缺少量化结果、与JD关键词不对齐）
                       - rewrite: 建议的改写示例（注意保持真实，不虚构经历）)

6. draft_resume       (字符串，针对“最终选中的 JD”生成的完整简历 Markdown 文本，
                       使用 # / ## 标题和 - 列表，突出与该 JD 相关的经历与成果，不要包含 JSON 转义字符)

【多 JD 匹配与选择】

7. target_jd_overview (数组，用于汇总每个候选 JD 的匹配情况。每个元素为对象：
                       - jd_index: 整数，和输入中的 JD 序号一致（从 1 开始）
                       - jd_title: 复制输入中 JD 标题（如：文件名或你看到的标题），不要自己造
                       - match_score: 0-100 整数，该 JD 与当前简历的匹配度
                       - recommendation_level: 字符串，如“强烈推荐”“可重点考虑”“可尝试”“不推荐”
                       - short_comment: 1-2 句专业点评，说明匹配好/不好的关键原因)

8. selected_jd_index  (整数，从 1 开始，表示你认为最适合做本轮深度优化的 JD 序号。
                       total_score / dimensions / draft_resume 等都应基于这个 JD。)

【学习与资源推荐】

9. learning_plan      (对象，字段：
                       - target_direction: 综合简历与 JD 后推荐的主要发展方向（如：AI产品、数据产品、数据分析等）
                       - summary: 用2-3句话概述未来3-6个月更理性、更有效的准备思路
                       - skills_to_focus: 数组，列出3-6个优先需要补齐或加强的能力/技能
                       - stages: 数组，每个阶段是一个对象，字段：
                           * name: 阶段名称，如“第1-4周：打基础”
                           * goals: 数组，该阶段的目标（站在求职者视角，而不是算法视角）
                           * actions: 数组，该阶段可以执行的具体行动建议（可操作，不要泛泛而谈）)

10. resources         (数组，每个元素是一个学习 / 面试资源建议对象，字段：
                       - platform: 平台名称，如“B站”“YouTube”“牛客”“CSDN”“其他”
                       - category: 资源类型，如“学习视频”“面试经验”“技术文章”“刷题/实战”
                       - search_keyword: 建议用户在该平台使用的搜索关键词（可以直接复制粘贴去搜）
                       - reason: 推荐理由，说明该资源如何帮助用户弥补当前简历中的短板或准备面试)

【相似岗位推荐（同方向）】

11. job_recommendations (数组，每个元素是一个岗位推荐对象，字段：
                       - title: 岗位名称，例如“AI 产品实习生”“数据产品实习生”
                       - company_type: 公司类型或示例描述，如“一线互联网公司（示例）”“数据智能公司（示例）”
                       - location: 城市或地区（可以是模糊描述，如“一线/新一线城市”）
                       - similarity_to_target_jd: 0-100整数，表示与「最终选中 JD」的相似程度
                       - match_reason: 推荐理由，说明为什么该岗位方向适合当前用户（结合简历与JD）
                       - core_requirements: 数组，列出3-6条该岗位核心要求示例（用自然语言）

强约束要求：
- 所有 job_recommendations 必须与「候选 JD 的岗位类型」同一职业族，例如：
  - 输入 JD 是 AI 产品 / 数据产品 / 互联网产品岗，只能推荐同类或高度相关产品/数据岗；
  - 不要跨到「财务、人力、纯后端开发」等完全不相关方向。
- 不要杜撰具体公司名和具体招聘链接，可以使用“某头部互联网公司（示例）”这类泛化描述。
- 所有内容必须基于【简历】和【候选 JD】的方向、技能差距来生成，避免和用户完全无关的建议。
- 不要编造简历中根本不存在的学校 / 公司 / 证书，可以合理推测适合的学习方向和资源关键词。
- 语气专业、友好，尽量站在求职者视角，避免空泛鸡汤，多给可执行建议。
- 输出必须是严格合法的 JSON，对象最外层必须包含上述所有字段。
    """

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": (
                        f"【简历文本】:\n{resume[:4000]}\n\n"
                        f"【候选岗位JD列表】（最多截取前2500字符/条）：\n\n{jd_combined}"
                    )
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"API 调用错误: {e}")
        return None


# ================= 4. UI 界面构建 =================

# --- Sidebar: 配置 ---
with st.sidebar:
    st.title("⚙️ 系统配置")

    config_mode = st.radio("运行模式", ["DeepSeek (推荐)", "OpenAI / 其他", "演示模式 (Demo)"])

    if config_mode == "DeepSeek (推荐)":
        st.info("💡 高性价比，逻辑能力强")
        api_key = st.text_input("DeepSeek API Key", type="password")
        base_url = "https://api.deepseek.com"
        model_name = "deepseek-chat"
    elif config_mode == "OpenAI / 其他":
        api_key = st.text_input("API Key", type="password")
        base_url = st.text_input("Base URL", value="https://api.openai.com/v1")
        model_name = st.text_input("Model Name", value="gpt-4o")
    else:
        api_key = "demo"
        base_url = ""
        model_name = "demo"

    st.markdown("---")
    st.markdown("### 使用指南")
    st.markdown("1. 上传简历 (PDF/Word/图片)")
    st.markdown("2. 粘贴或上传 1–N 个 JD（文本/文件/图片）")
    st.markdown("3. 点击分析")
    st.markdown("4. 查看匹配报告 + 学习建议 + 岗位推荐并下载简历")

# --- Main Area ---
st.title("💼 JobAlign AI Pro | 职配助手")
st.caption("多岗位匹配 + 简历优化 + 学习规划 + 岗位推荐，一次走完。")

col1, col2 = st.columns(2)

# ========= 4.1 简历输入 =========
with col1:
    st.subheader("1. 个人简历")
    resume_file = st.file_uploader(
        "上传简历（支持 PDF / Word / 文本 / 图片）",
        type=['pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif']
    )
    resume_text = ""
    if resume_file:
        resume_text = DocumentHandler.extract_text(resume_file)
        if resume_text.startswith("Error: 文件解析失败"):
            st.error(resume_text)
        else:
            st.success(f"✅ 已提取约 {len(resume_text)} 字")
            with st.expander("查看简历解析内容"):
                st.text(resume_text[:800] + "..." if len(resume_text) > 800 else resume_text)
    else:
        resume_text = st.text_area("或直接粘贴简历内容", height=200)

# ========= 4.2 多 JD 输入 =========
with col2:
    st.subheader("2. 目标岗位 (JD) — 可一次上传多个")
    jd_input_method = st.radio("输入方式", ["文本粘贴（单个）", "文件上传（可多个）"], horizontal=True)
    jd_entries = []

    if jd_input_method == "文本粘贴（单个）":
        jd_text = st.text_area(
            "粘贴职位描述（单个 JD）",
            height=220,
            placeholder="职位描述\n岗位职责...\n任职要求..."
        )
        if jd_text.strip():
            jd_entries.append({
                "index": 1,
                "title": "文本JD",
                "text": jd_text
            })
    else:
        jd_files = st.file_uploader(
            "上传 JD 文件（可多选，支持 PDF / Word / 文本 / 图片）",
            type=['pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif'],
            accept_multiple_files=True
        )
        if jd_files:
            for idx, jf in enumerate(jd_files, start=1):
                text = DocumentHandler.extract_text(jf)
                if text.startswith("Error: 文件解析失败"):
                    st.error(f"❌ JD 文件解析失败：{jf.name}，请检查后重试。")
                    continue
                jd_entries.append({
                    "index": idx,
                    "title": jf.name,
                    "text": text
                })
            if jd_entries:
                st.success(f"✅ 已成功导入 {len(jd_entries)} 个 JD")
                with st.expander("查看部分 JD 内容预览"):
                    for entry in jd_entries:
                        st.markdown(f"**[{entry['index']}] {entry['title']}**")
                        preview = entry['text']
                        st.text(preview[:400] + "..." if len(preview) > 400 else preview)
                        st.markdown("<hr style='margin: 4px 0; opacity: 0.3'/>", unsafe_allow_html=True)

st.markdown("---")

# 提交按钮
btn_col1, btn_col2, btn_col3 = st.columns([1, 2, 1])
with btn_col2:
    analyze_btn = st.button(
        "🚀 开始深度分析：多 JD 匹配 + 简历优化 + 学习建议 + 岗位推荐",
        use_container_width=True
    )

# 逻辑处理
if analyze_btn:
    if not resume_text or not resume_text.strip():
        st.warning("⚠️ 请先上传或粘贴简历。")
    elif not jd_entries:
        st.warning("⚠️ 请至少提供 1 个 JD（可多选）。")
    elif config_mode != "演示模式 (Demo)" and not api_key:
        st.error("⚠️ 请输入 API Key 才能使用 AI 功能。")
    else:
        with st.spinner("🤖 AI 正在阅读你的简历 & 多个 JD，并生成匹配报告与成长建议..."):
            if config_mode == "演示模式 (Demo)":
                time.sleep(2)
                result = MOCK_DATA
            else:
                result = analyze_with_llm(api_key, base_url, model_name, resume_text, jd_entries)

            if result:
                st.session_state.result_json = result
                st.session_state.analyzed = True
                st.rerun()

# ================= 5. 结果展示 =================
if st.session_state.analyzed and st.session_state.result_json:
    res = st.session_state.result_json

    # ----- 5.1 多 JD 匹配概览 -----
    st.header("📌 多岗位匹配概览")
    jd_overview = res.get("target_jd_overview", [])
    selected_jd_index = res.get("selected_jd_index", None)

    if jd_overview:
        df_jd = pd.DataFrame([
            {
                "序号": item.get("jd_index"),
                "岗位名称": item.get("jd_title"),
                "匹配分": item.get("match_score"),
                "推荐级别": item.get("recommendation_level"),
                "点评": item.get("short_comment")
            }
            for item in jd_overview
        ])
        st.dataframe(df_jd, use_container_width=True)

        if selected_jd_index:
            selected_row = next(
                (item for item in jd_overview if item.get("jd_index") == selected_jd_index),
                None
            )
            if selected_row:
                st.success(
                    f"本轮详细优化基于：第 {selected_jd_index} 个岗位 —— {selected_row.get('jd_title', '')}"
                )
    else:
        st.info("暂无多 JD 匹配概览数据。")

    st.markdown("---")

    # ----- 5.2 匹配分 & 亮点 / 缺失 -----
    st.header("📊 针对选中 JD 的匹配报告")
    m_col1, m_col2 = st.columns([1, 1])

    with m_col1:
        st.metric("总体匹配得分", res.get('total_score', 0), delta_color="normal")
        # 雷达图
        dimensions = res.get('dimensions', {})
        if dimensions:
            df_radar = pd.DataFrame(dict(
                r=list(dimensions.values()),
                theta=list(dimensions.keys())
            ))
            fig = px.line_polar(df_radar, r='r', theta='theta', line_close=True, range_r=[0, 100])
            fig.update_traces(fill='toself')
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("暂无维度评分数据。")

    with m_col2:
        st.subheader("🎯 核心发现")
        tab_high, tab_gap = st.tabs(["✨ 亮点 (Highlights)", "⚠️ 缺失 / 风险 (Gaps)"])
        with tab_high:
            for i in res.get('highlights', []):
                st.success(f"• {i}")
        with tab_gap:
            for i in res.get('gaps', []):
                st.error(f"• {i}")

    st.markdown("---")

    # ----- 5.3 智能改写建议 -----
    st.subheader("💡 智能改写建议（逐条对比）")
    suggestions = res.get('suggestions', [])
    if suggestions:
        for item in suggestions:
            with st.container():
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown(f"**🔴 原文 ({item.get('section', '未标注模块')})**")
                    st.caption(f"问题：{item.get('problem', '未提供问题说明')}")
                    st.text(item.get('original', ''))
                with c2:
                    st.markdown("**🟢 优化后示例**")
                    st.info(item.get('rewrite', ''))
                st.markdown("<hr style='margin: 5px 0; opacity: 0.3'/>", unsafe_allow_html=True)
    else:
        st.info("暂无改写建议。")

    st.markdown("---")

    # ----- 5.4 相似岗位推荐 -----
    job_recs = res.get("job_recommendations", [])
    if job_recs:
        st.header("🔍 相关岗位推荐（同方向）")
        st.caption("以下为同一职业方向下的示例岗位画像，方便你拓展可投递的公司与职位方向。")
        for job in job_recs:
            with st.container():
                title = job.get("title", "未知岗位")
                company_type = job.get("company_type", "")
                location = job.get("location", "")
                similarity = job.get("similarity_to_target_jd", None)
                match_reason = job.get("match_reason", "")
                core_reqs = job.get("core_requirements", [])

                st.markdown(f"**{title}**")
                meta = []
                if company_type:
                    meta.append(company_type)
                if location:
                    meta.append(location)
                if isinstance(similarity, (int, float)):
                    meta.append(f"与当前目标 JD 相似度约 {similarity} 分")
                if meta:
                    st.caption(" · ".join(meta))
                if match_reason:
                    st.write(match_reason)
                if core_reqs:
                    st.markdown("核心要求示例：")
                    for r_item in core_reqs:
                        st.write(f"- {r_item}")
                st.markdown("<hr style='margin: 5px 0; opacity: 0.15'/>", unsafe_allow_html=True)

    st.markdown("---")

    # ----- 5.5 学习与成长建议 -----
    learning_plan = res.get("learning_plan")
    if learning_plan:
        st.header("📚 学习与成长建议（未来 3–6 个月参考）")
        st.subheader(f"推荐发展方向：{learning_plan.get('target_direction', '未识别')}")
        summary = learning_plan.get('summary')
        if summary:
            st.write(summary)

        skills_to_focus = learning_plan.get('skills_to_focus', [])
        if skills_to_focus:
            st.markdown("**优先关注的能力 / 技能：**")
            for s in skills_to_focus:
                st.write(f"- {s}")

        stages = learning_plan.get('stages', [])
        if stages:
            st.markdown("**阶段性行动建议：**")
            for stage in stages:
                with st.expander(stage.get('name', '未命名阶段'), expanded=False):
                    goals = stage.get('goals', [])
                    actions = stage.get('actions', [])
                    if goals:
                        st.markdown("📌 阶段目标：")
                        for g in goals:
                            st.write(f"- {g}")
                    if actions:
                        st.markdown("🧭 推荐行动：")
                        for a in actions:
                            st.write(f"- {a}")
    else:
        st.info("暂无学习规划数据。")

    st.markdown("---")

    # ----- 5.6 学习资源 & 面经推荐 -----
    resources = res.get("resources", [])
    if resources:
        st.header("🎥 学习资源 & 面试经验推荐")
        st.caption("以下为「平台 + 搜索关键词」形式，建议复制关键词到对应平台搜索最新内容。")
        for r in resources:
            with st.container():
                platform = r.get('platform', '其他')
                category = r.get('category', '')
                search_keyword = r.get('search_keyword', '')
                reason = r.get('reason', '')

                st.markdown(f"**{platform} · {category}**")
                if search_keyword:
                    st.markdown("推荐搜索关键词：")
                    st.code(search_keyword, language="text")
                if reason:
                    st.caption(reason)
                st.markdown("<hr style='margin: 5px 0; opacity: 0.2'/>", unsafe_allow_html=True)
    else:
        st.info("暂无资源推荐数据。")

    st.markdown("---")

    # ----- 5.7 简历生成与导出 -----
    st.header("📝 定制版简历预览与导出")

    draft_resume = res.get('draft_resume', '')
    if draft_resume:
        # 生成 Word 文档
        docx_file = WordGenerator.create_docx_from_markdown(draft_resume)

        # 导出按钮
        st.download_button(
            label="📥 下载 Word (.docx) 简历",
            data=docx_file,
            file_name=f"JobAlign_定制简历.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        with st.expander("查看 / 编辑简历 Markdown 源码"):
            st.text_area("简历 Markdown 源码", value=draft_resume, height=400)
    else:
        st.info("暂无定制简历内容。")
